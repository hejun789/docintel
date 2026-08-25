"""Tests for chunk_text — the heart of the ingestion pipeline.

A silent bug here (wrong metadata, dropped chunks) corrupts everything
downstream, so chunk count, metadata, and indexing are worth pinning down.
"""
from ingest import chunk_text, extract_text_from_docx


def test_short_text_produces_one_chunk():
    pages = [{"text": "A short sentence about phishing.", "page_number": 1}]
    chunks = chunk_text(pages, "doc.pdf")
    assert len(chunks) == 1


def test_long_text_splits_into_multiple_chunks():
    # ~2000 chars, well over the 512-char chunk size, so it must split.
    pages = [{"text": "Phishing detection with large language models. " * 50, "page_number": 1}]
    chunks = chunk_text(pages, "doc.pdf")
    assert len(chunks) > 1


def test_every_chunk_has_required_metadata():
    pages = [{"text": "Some content here. " * 60, "page_number": 7}]
    chunks = chunk_text(pages, "report.pdf")
    for c in chunks:
        assert set(c) >= {"text", "source", "page_number", "chunk_index"}
        assert c["source"] == "report.pdf"
        assert c["page_number"] == 7
        assert c["text"].strip()


def test_source_is_basename_only():
    pages = [{"text": "Body text.", "page_number": 1}]
    chunks = chunk_text(pages, "/some/nested/path/report.pdf")
    assert chunks[0]["source"] == "report.pdf"


def test_chunk_index_is_sequential_across_pages():
    pages = [
        {"text": "First page content. " * 60, "page_number": 1},
        {"text": "Second page content. " * 60, "page_number": 2},
    ]
    chunks = chunk_text(pages, "doc.pdf")
    indices = [c["chunk_index"] for c in chunks]
    assert indices == list(range(len(chunks)))


def test_empty_pages_produce_no_chunks():
    assert chunk_text([], "doc.pdf") == []


def test_docx_extraction_includes_table_text(tmp_path):
    import docx as python_docx
    doc = python_docx.Document()
    doc.add_paragraph("Assignment brief intro.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Distinction"
    table.rows[0].cells[1].text = "70% and above"
    path = tmp_path / "rubric.docx"
    doc.save(str(path))

    text = " ".join(p["text"] for p in extract_text_from_docx(str(path)))
    assert "Assignment brief intro." in text
    assert "Distinction" in text          # table content must be captured
    assert "70% and above" in text


def test_chunking_uses_config_values(monkeypatch):
    """CHUNK_SIZE/CHUNK_OVERLAP in config must actually drive the splitter."""
    import ingest
    monkeypatch.setattr(ingest, "CHUNK_SIZE", 100)
    monkeypatch.setattr(ingest, "CHUNK_OVERLAP", 0)
    pages = [{"text": "word " * 200, "page_number": 1}]
    chunks = ingest.chunk_text(pages, "doc.pdf")
    assert all(len(c["text"]) <= 100 for c in chunks)
