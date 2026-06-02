import os
import fitz  # PyMuPDF
import docx as python_docx
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document as LCDocument
from config import CHROMA_DIR, CHROMA_COLLECTION, EMBEDDING_MODEL

_embedder = None
_vectorstore = None


def _get_embedder():
    global _embedder
    if _embedder is None:
        _embedder = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
    return _embedder


def _get_vectorstore():
    global _vectorstore
    if _vectorstore is None:
        _vectorstore = Chroma(
            collection_name=CHROMA_COLLECTION,
            embedding_function=_get_embedder(),
            persist_directory=CHROMA_DIR,
        )
    return _vectorstore


def extract_text_from_pdf(filepath: str) -> list[dict]:
    """
    Extract text from each page of a PDF.
    Returns a list of dicts: {'text': str, 'page_number': int}
    """
    pages = []

    try:
        doc = fitz.open(filepath)
    except Exception as e:
        raise ValueError(f"Could not open PDF '{filepath}': {e}")

    for page_index in range(len(doc)):
        page = doc[page_index]
        text = page.get_text("text")  # plain text, preserves reading order

        # Skip pages that are blank or nearly blank
        cleaned = text.strip()
        if not cleaned:
            continue

        pages.append({
            "text": cleaned,
            "page_number": page_index + 1  # 1-based for human readability
        })

    doc.close()
    return pages


def extract_text_from_docx(filepath: str) -> list[dict]:
    doc = python_docx.Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    pages = []
    page_size = 15  # group paragraphs into pseudo-pages
    for i in range(0, len(paragraphs), page_size):
        text = "\n".join(paragraphs[i:i + page_size])
        pages.append({"text": text, "page_number": (i // page_size) + 1})
    return pages


def extract_text_from_txt(filepath: str) -> list[dict]:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        lines = [l.rstrip() for l in f.readlines()]
    pages = []
    page_size = 50  # group lines into pseudo-pages
    for i in range(0, len(lines), page_size):
        text = "\n".join(lines[i:i + page_size]).strip()
        if text:
            pages.append({"text": text, "page_number": (i // page_size) + 1})
    return pages


def extract_text_from_file(filepath: str) -> list[dict]:
    ext = filepath.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(filepath)
    elif ext == "docx":
        return extract_text_from_docx(filepath)
    elif ext in ("txt", "md"):
        return extract_text_from_txt(filepath)
    raise ValueError(f"Unsupported file type: .{ext}")


def chunk_text(pages: list[dict], filename: str) -> list[dict]:
    """
    Split page text into overlapping chunks, each carrying source metadata.
    Returns a list of dicts: {'text', 'source', 'page_number', 'chunk_index'}
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=512,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    chunks = []
    chunk_index = 0
    source = os.path.basename(filename)

    for page in pages:
        page_chunks = splitter.split_text(page["text"])
        for chunk_text_content in page_chunks:
            chunks.append({
                "text": chunk_text_content,
                "source": source,
                "page_number": page["page_number"],
                "chunk_index": chunk_index,
            })
            chunk_index += 1

    return chunks


def store_chunks(chunks: list[dict]) -> int:
    """
    Embed each chunk and store it in ChromaDB via LangChain's Chroma vectorstore.
    Deletes existing chunks for the same source first (upsert behaviour).
    Returns the number of chunks stored.
    """
    vectorstore = _get_vectorstore()

    # Remove any existing chunks for this source before re-ingesting
    source = chunks[0]["source"] if chunks else None
    if source:
        existing = vectorstore._collection.get(where={"source": source})
        if existing["ids"]:
            vectorstore.delete(ids=existing["ids"])

    documents = [
        LCDocument(
            page_content=c["text"],
            metadata={
                "source": c["source"],
                "page_number": c["page_number"],
                "chunk_index": c["chunk_index"],
            },
        )
        for c in chunks
    ]
    ids = [f"{c['source']}_chunk_{c['chunk_index']}" for c in chunks]
    vectorstore.add_documents(documents, ids=ids)
    return len(chunks)


def list_documents() -> list[str]:
    """Return sorted list of unique source filenames stored in ChromaDB."""
    vectorstore = _get_vectorstore()
    results = vectorstore._collection.get(include=["metadatas"])
    if not results["metadatas"]:
        return []
    return sorted({m["source"] for m in results["metadatas"] if m})


def delete_document(filename: str) -> int:
    """
    Remove all chunks belonging to a document from ChromaDB.
    Returns the number of chunks deleted.
    """
    vectorstore = _get_vectorstore()
    results = vectorstore._collection.get(where={"source": filename})
    count = len(results["ids"])
    if count > 0:
        vectorstore.delete(ids=results["ids"])
    return count


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python ingest.py <path_to_pdf>")
        sys.exit(1)

    filepath = sys.argv[1]

    print("Step 1: Extracting text...")
    pages = extract_text_from_pdf(filepath)
    print(f"  {len(pages)} non-blank pages\n")

    print("Step 2: Chunking...")
    chunks = chunk_text(pages, filepath)
    print(f"  {len(chunks)} chunks\n")

    print("Step 3: Embedding and storing...")
    stored = store_chunks(chunks)
    print(f"  {stored} chunks stored in ChromaDB\n")

    print("Verifying — querying ChromaDB directly:")
    from sentence_transformers import SentenceTransformer
    from config import EMBEDDING_MODEL
    embedder = SentenceTransformer(EMBEDDING_MODEL)
    collection = _get_collection()
    test_query = "network security topology"
    query_embedding = embedder.encode([test_query]).tolist()
    results = collection.query(query_embeddings=query_embedding, n_results=3)
    for i, doc in enumerate(results["documents"][0]):
        meta = results["metadatas"][0][i]
        print(f"  Result {i+1} | Page {meta['page_number']} | {meta['source']}")
        print(f"  {doc[:200]}\n")
